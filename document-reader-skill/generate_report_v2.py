"""Generate restructured vertical mill report — Jinchuan standard."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '仿宋'
font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
style.paragraph_format.line_spacing = Pt(28)
style.paragraph_format.space_after = Pt(0)

# ── 辅助函数 ──
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        sz = {1: 16, 2: 15, 3: 14}.get(level, 14)
        run.font.size = Pt(sz)
    return h

def para(text, indent=True, bold_head=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '仿宋'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(14)
    if bold_head:
        run.bold = True
    return p

def para_split(head, body, indent=True):
    """写一段话，开头几个字加粗作为标志词。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r1 = p.add_run(head)
    r1.bold = True
    r1.font.name = '仿宋'
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    r1.font.size = Pt(14)
    r2 = p.add_run(body)
    r2.font.name = '仿宋'
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    r2.font.size = Pt(14)
    return p

def blank():
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
run = title.add_run('印尼金川WP公司立式磨存在问题及整改措施')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run('冶 炼 部 设 备 管 理')
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(14)

# ════════════════════════════════════════
# 一、基本情况
# ════════════════════════════════════════
heading('一、基本情况')

para(
    '公司粉煤制备系统采用布莱克散料处理技术（北京）有限公司生产的BRM28.3M型立式辊磨机，'
    '设计产能55t/h（干粉），系统设计风量240,000m³/h，自2019年投运以来已连续运行六年，'
    '承担冶炼系统煤粉制备任务。随着设备运行年限增加，各主要部件逐步出现不同程度的磨损与老化。',
)
para(
    '近期结合设备运行数据及现场检查，对立式磨系统进行了全面排查，发现磨盘衬板、选粉机、'
    '布袋收尘器、磨辊润滑密封系统及高挥发份煤种应用安全等方面存在一定问题。'
    '为提升设备运行可靠性，保障系统安全稳定运行，现对立式磨系统存在的主要问题进行分析，'
    '并制定相应整改措施，结合生产组织分步推进实施。',
)

# ════════════════════════════════════════
# 二、存在的主要问题
# ════════════════════════════════════════
heading('二、存在的主要问题')

# ---- (一) ----
heading('（一）磨盘衬板磨损严重', level=2)
para(
    '磨盘衬板自设备投运以来持续运行至今，长期受到物料冲刷与研磨作用，工作面磨损较为明显，'
    '部分区域已接近或超过设计磨损极限。衬板磨损后，将导致研磨效率下降、料层稳定性变差、'
    '系统电耗增加；若长期超限运行，可能进一步影响磨盘本体结构安全及设备中长期稳定运行。'
    '目前该问题已到必须安排更换的阶段。',
)

# ---- (二) ----
heading('（二）选粉机动叶片磨损', level=2)
para(
    '选粉机作为煤粉细度控制的核心设备，其动叶片长期处于高速旋转及含尘气流冲刷工况下，'
    '迎风面磨损明显，导致分级效率下降。当前选粉机运行频率持续维持47Hz，高于正常带载运行区间的'
    '30-40Hz，反映出因叶片磨损后分级效能下降，系统被迫以更高转速运行方能达到细度要求。'
    '若磨损持续发展，将引起煤粉细度波动，影响回转窑燃烧工况，严重时可能发生叶片脱落，'
    '造成选粉机本体损坏及计划外停产。选粉机整体备件已在库，可满足整体更换需求。',
)

# ---- (三) ----
heading('（三）布袋收尘器布袋老化', level=2)
para(
    '布袋收尘器滤袋自2021年12月更换至今已连续运行近四年，接近常规使用寿命。'
    '聚酯针刺毡滤袋长期处于含煤粉烟气、脉冲喷吹循环及温度波动工况，存在滤袋纤维疲劳、'
    '强度下降、透气性恶化等老化趋势。布袋性能下降后将直接导致系统通风阻力增大、'
    '引风机电耗上升，若局部破损则可能造成粉尘排放超标，不符合日益严格的环保要求。'
    '布袋备件已在库，但存放已近四年，需对现有库存状态进行核查。',
)

# ---- (四) ----
heading('（四）磨辊润滑密封系统存在隐患', level=2)

heading('1.#2磨辊密封磨损', level=3)
para(
    '#2磨辊密封自投产至今一直未更换，密封件经过长期运行已严重磨损，煤粉颗粒已渗入润滑系统，'
    '回油中可见煤粉。目前进回油量及润滑功能尚基本正常，但随着密封失效持续加重，煤粉将加速侵入，'
    '最终导致润滑系统损坏及磨辊轴承烧毁，属于典型的渐进式设备隐患。需结合近期停产窗口安排密封件更换。',
)

heading('2.#1磨辊密封风不足', level=3)
para(
    '#1磨辊位于磨辊密封风管道的末端，由于管路距离远、沿程阻力损失大，实际到达密封罩的风压和风量'
    '明显不足，无法在密封罩内形成有效的正压气幕。煤粉由此进入密封罩内，逐步磨损密封件并威胁润滑系统。'
    '该问题属于系统管路布置固有的缺陷，需通过技术改造加以解决。目前拟先采取单独加装密封风管路的方案，'
    '若效果不理想再研究更换密封方式。',
)

# ---- (五) ----
heading('（五）高挥发份煤种应用安全保障需进一步完善', level=2)
para(
    '随着高挥发份煤种采购比例逐步提高，制粉系统的防火防爆及运行控制面临更高要求。'
    '相关技术规范中明确要求：采用烟煤时制粉系统气粉混合物中氧含量应控制≤14%，'
    '采用褐煤时氧含量应控制≤12%。目前系统烟气含氧量指标按≤12%从严控制。'
    '但在实际运行中，系统仍存在一定漏风量、热风炉有产生火花的潜在风险，'
    '且高挥发份煤常态化使用后制粉系统安全裕度相对下降，'
    '现有安全保障及监测联锁措施仍有进一步完善的空间。',
)

# ════════════════════════════════════════
# 三、整改措施
# ════════════════════════════════════════
heading('三、整改措施')

heading('（一）利用停产窗口实施重点检修', level=2)
para(
    '结合各生产线停产间隙，优先安排必须停机方可实施的检修项目。',
    indent=False,
)
para(
    '磨盘衬板更换：先利用7月15日左右4#线升温、3#线停产期间，组织对磨盘衬板进行试拆装，'
    '评估螺栓锈蚀程度、拆卸难度及所需专用工具，形成书面评估报告。在确认拆除方案可行后，'
    '于3#线升温、1#线停产窗口期正式实施衬板更换。更换前逐件核对备件型号与材质'
    '（图号BRM28.3M.9），更换后空载试运行不少于30分钟确认无异常后投入正常使用。',
)
para(
    '#2磨辊密封更换：与磨盘衬板更换同步安排实施，减少单独停机次数。'
    '更换前加强润滑油定期取样检测（每两周一次），掌握磨损发展趋势确定更换时机。'
    '更换密封时同步清洗润滑管路及油箱、更换新油，确保系统内无残留煤粉。'
    '密封安装严格按BRM28.3M立式辊磨机安装手册相关技术要求执行。',
)

heading('（二）结合年度大修推进系统更新', level=2)
para(
    '对可在年度检修中统一规划的系统性更新项目，纳入明年大修计划统筹安排。',
    indent=False,
)
para(
    '选粉机更新：选粉机整体备件已在库，整体更换纳入明年度大修计划，与动氧系统年度检修同步实施，'
    '充分利用停产窗口。在大修前开展叶片更换及现场动平衡可行性的技术研究，探索不整体拆机条件下的'
    '快速更换方案，为后续缩短检修时间积累技术储备。同时将当前选粉机运行电流、振动值、转速及'
    '煤粉细度纳入每周例行监测，做好趋势分析。',
)
para(
    '布袋收尘器布袋更换：立即组织对各仓室布袋进行抽样检查，每仓不少于三条，重点检查袋口、'
    '底部折边及与笼骨接触部位的磨损与老化情况，形成检查台账。根据抽检结果评估各仓布袋整体状况，'
    '制定分仓更换计划。同步检查库存储备布袋是否存在受潮、发霉或老化脆变等问题，'
    '如已严重老化需重新采购。更换完成后做好喷吹参数调试，建立布袋运行寿命档案。',
)

heading('（三）持续完善运行保障措施', level=2)
para(
    '对不需要长时间停产即可实施的改进项目，纳入日常运维持续落实。',
    indent=False,
)
para(
    '密封风改造：为#1磨辊单独加装一路密封风管路，直接从密封风机出口引支管至#1磨辊密封罩，'
    '确保风压风量满足设计要求。加装后连续跟踪不少于一个月，定期检查润滑油质变化，'
    '评估改造效果。若单独加装方案效果有限，再研究采用接触式密封或组合式密封的改造方案。',
)
para(
    '火花捕捉装置加装：在热风炉出口至磨机入口之间的管道上加装火花补给器（火花捕捉器），'
    '作为防止明火进入磨机系统的物理屏障。选型应满足烟气温度不低于300℃的耐温要求，'
    '捕捉效率不低于95%，并配备差压检测及报警装置以便及时清理。',
)
para(
    '系统漏风治理：全面排查制粉系统各法兰连接处、人孔门、检查孔、锁风阀等潜在漏风点并实施密封处理，'
    '在保证正常通风需求的前提下尽可能降低漏风率，抑制外部空气（氧含量21%）进入系统导致含氧量上升。',
)
para(
    '在线监测与联锁完善：加强氧含量分析仪校准维护（每月至少一次），完善CO浓度监测与磨机联锁保护逻辑，'
    'CO≥800PPM时报警、≥1200PPM时联锁停机并充入惰性气体。优化现有氮气惰化方案，'
    '确保紧急状态下能快速向磨机、布袋收尘器和粉煤仓注入足量氮气。',
)
para(
    '运行管理标准化：编制高挥发份煤种制粉操作指导书，明确不同煤种配比下的氧含量、'
    '温度控制范围及加料速率调整原则，操作人员经专项培训后方可上岗。'
    '建立高挥发份煤种运行台账，每日记录煤种配比、氧含量、温度、CO等关键参数，定期分析趋势及时调整。',
)

# ════════════════════════════════════════
# 四、整改计划
# ════════════════════════════════════════
heading('四、整改计划')
blank()

table = doc.add_table(rows=6, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['整改项目', '整改目标', '计划时间', '责任单位']
data = [
    ['磨盘衬板更换', '恢复研磨效率\n消除本体安全隐患', '试拆装：7月中旬\n正式更换：结合1#线\n停产窗口',
     '冶炼部\n生产保障部'],
    ['#2磨辊密封更换', '恢复密封性能\n保障润滑系统安全', '结合磨盘衬板更换\n同步实施',
     '冶炼部\n生产保障部'],
    ['选粉机更新', '恢复分级效率\n降低运行转速', '纳入明年年度大修\n统筹安排',
     '冶炼部\n设备管理'],
    ['布袋收尘器\n布袋更换', '恢复通风性能\n确保达标排放', '抽检评估：立即开展\n（2周内完成）\n分仓更换：逐步实施',
     '冶炼部\n设备管理'],
    ['高挥发份煤种\n安全保障措施', '提升系统防爆能力\n完善监测联锁', '火花捕捉器：3个月内\n完成安装\n其他措施：持续完善',
     '冶炼部\n设备管理\n供应部'],
]

# 表头
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(11)
    set_cell_shading(cell, 'D9E2F3')

# 数据行
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '仿宋'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                run.font.size = Pt(11)

# 设置列宽
from docx.shared import Cm
widths = [Cm(3.0), Cm(4.5), Cm(4.0), Cm(3.0)]
for row in table.rows:
    for j, cell in enumerate(row.cells):
        cell.width = widths[j]

blank()

# ════════════════════════════════════════
# 五、下一步工作
# ════════════════════════════════════════
heading('五、下一步工作')

para(
    '下一步，冶炼部将按照"统筹安排、分步实施、风险可控"的原则，结合各生产线停产窗口及年度检修计划，'
    '有序推进各项整改措施落实，确保设备检修与生产组织协调推进，最大限度减少对冶炼系统的影响。',
)
para(
    '同时，将进一步加强设备状态监测、运行分析和隐患排查，不断完善设备全生命周期管理机制，'
    '持续提升立式磨系统安全、稳定、经济运行水平，为冶炼系统连续稳定生产提供可靠保障。',
)

# ── 落款 ──
blank()
blank()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run('印尼金川WP公司冶炼部')
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(14)
run.bold = True

blank()
p = doc.add_paragraph()
run = p.add_run('2025年7月')
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(14)

# ── 保存 ──
output = r'D:\工作\设备资料\分类资料\立式磨\印尼金川WP公司立式磨存在问题及整改措施报告.docx'
doc.save(output)
print(f'OK: {output}')
