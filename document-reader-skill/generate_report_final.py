"""Generate the vertical mill rectification report — FINAL submission version."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ═══════════════════════════════════════════
# 页面设置
# ═══════════════════════════════════════════
for section in doc.sections:
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

# ═══════════════════════════════════════════
# 样式
# ═══════════════════════════════════════════
style_normal = doc.styles['Normal']
style_normal.font.name = '仿宋'
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
style_normal.font.size = Pt(16)
style_normal.paragraph_format.line_spacing = Pt(28)
style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)

def sf(run, name='仿宋', size=16, bold=False):
    run.font.name = name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(32)
    run = p.add_run(text)
    sf(run, '黑体', 22, True)
    return p

def add_doc_number(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(28)
    run = p.add_run(text)
    sf(run, '仿宋', 16)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    sf(run, '黑体', 16, True)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    sf(run, '楷体', 16, True)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    sf(run, '仿宋', 16, True)
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    sf(run, '仿宋', 16)
    return p

def add_body_ni(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    sf(run, '仿宋', 16)
    return p

def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_after = Pt(0)
    return p

def add_right(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run(text)
    sf(run, '仿宋', 16)
    return p

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell(cell, text, name='仿宋', size=12, bold=False, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    sf(run, name, size, bold)

# ═══════════════════════════════════════════
# 正文
# ═══════════════════════════════════════════

add_doc_number('印尼金川WP公司冶炼部〔2026〕XX号')
add_title('关于立式磨存在问题及整改措施的报告')
add_body_ni('公司领导：')

add_body(
    '公司粉煤制备系统采用布莱克散料处理技术（北京）有限公司生产的BRM28.3M型立式辊磨机，'
    '设计产能55t/h（干粉），自2019年投运以来已连续运行六年，承担冶炼系统煤粉制备任务。'
    '随着设备运行年限增加，各主要部件逐步出现不同程度的磨损与老化。'
)
add_body(
    '近期结合设备运行数据及现场检查，对立式磨系统进行了全面排查，发现磨盘衬板、选粉机、'
    '布袋收尘器、磨辊润滑密封系统及高挥发份煤种应用安全等方面存在一定问题。'
    '为保障设备安全稳定运行，现将有关情况及整改措施报告如下。'
)

# ═══════════════════════════════════════════
# 一、存在的主要问题
# ═══════════════════════════════════════════
add_h1('一、存在的主要问题')

# ——（一）磨盘衬板 ——
add_h2('（一）磨盘衬板磨损较大')
add_body(
    '磨盘衬板自设备投运以来持续运行至今，长期受到物料冲刷与研磨作用，工作面磨损较为明显，'
    '部分区域已接近设计磨损极限。衬板磨损后，将导致研磨效率下降、料层稳定性变差、'
    '系统电耗增加，若继续超限运行，可能对磨盘本体结构长期安全产生一定影响。'
)
add_body(
    '综合评估认为，该问题已具备更换条件，建议结合停产窗口组织实施。'
)

# ——（二）选粉机 ——
add_h2('（二）选粉机转子动平衡失衡及动叶片磨损')
add_body(
    '选粉机转子在长期运行过程中存在动平衡失衡问题，高速旋转时偏心振动较大，'
    '为保证设备安全，目前运行频率控制在30Hz左右，设备能力未能充分发挥。'
)
add_body(
    '动叶片长期受到含尘气流冲刷，迎风面磨损较为明显，导致分级效率下降，'
    '煤粉细度控制稳定性降低。'
)
add_body(
    '综合来看，目前设备仍可运行，但已对运行效率和煤粉质量控制产生一定影响，'
    '建议结合年度大修实施整体更换，并同步解决转子动平衡问题。'
)

# ——（三）布袋收尘器 ——
add_h2('（三）布袋收尘器滤袋老化')
add_body(
    '布袋收尘器滤袋自2021年12月更换至今已连续运行近四年。目前系统运行压差总体正常，'
    '但随着运行时间增加，滤袋纤维疲劳、强度下降、透气性恶化等老化风险逐步增大。'
    '滤袋性能下降后将导致系统通风阻力增大、引风机电耗上升，若局部破损还可能造成粉尘排放超标。'
)
add_body(
    '综合评估认为，该问题应结合滤袋抽检结果，合理安排分仓或整体更换。'
)

# ——（四）磨辊密封 ——
add_h2('（四）磨辊润滑密封系统存在不足')
add_h3('1. #2磨辊密封磨损')
add_body(
    '#2磨辊密封自投产至今一直未更换，密封件经长期运行已出现磨损，煤粉颗粒已渗入润滑系统，'
    '回油中可见煤粉。目前进回油量及润滑功能尚基本正常，但随着密封磨损持续加重，'
    '煤粉将持续进入润滑系统，加剧轴承磨损，对磨辊长期可靠运行产生不利影响。'
)
add_body(
    '综合评估认为，该问题建议结合停产窗口与磨盘衬板更换同步实施。'
)
add_h3('2. #1磨辊密封风不足')
add_body(
    '#1磨辊位于密封风管道末端，由于管路距离远、沿程阻力损失较大，实际到达密封罩的风压风量'
    '不足，难以在密封罩内形成有效的正压气幕，煤粉由此进入密封罩内逐步磨损密封件，'
    '对润滑系统长期稳定运行产生不利影响。该问题属于管路布置固有缺陷，需通过技术改造加以解决。'
)
add_body(
    '综合评估认为，该问题建议先采取单独加装密封风管路的方案，若效果不理想再研究更换密封方式。'
)

# ——（五）高挥发份煤 ——
add_h2('（五）高挥发份煤种应用安全保障需进一步完善')
add_body(
    '随着高挥发份煤种采购比例逐步提高，对制粉系统的防火防爆管理提出了更高要求。'
    '相关技术规范要求：采用烟煤时制粉系统气粉混合物中氧含量应控制≤14%，'
    '采用褐煤时氧含量应控制≤12%。目前系统烟气含氧量指标按≤12%从严控制。'
    '但在实际运行中，系统仍存在一定漏风量，热风炉有产生火花的潜在风险，'
    '现有漏风控制、火源控制及监测联锁等措施仍有进一步完善的空间。'
)
add_body(
    '综合评估认为，该问题应纳入常态化运行保障范围，通过技术改造和运行管理持续完善。'
)

# ═══════════════════════════════════════════
# 二、整改措施
# ═══════════════════════════════════════════
add_h1('二、整改措施')

add_body(
    '针对上述问题，按照"轻重缓急、分类实施、统筹检修"的原则，'
    '按以下三类措施组织实施。'
)

# ——（一）停产窗口 ——
add_h2('（一）利用停产窗口实施重点检修')
add_body_ni('整改目的：对影响设备安全运行且需停机实施的项目，结合停产窗口集中组织检修。')

add_h3('1. 磨盘衬板更换')
add_body(
    '主要措施：先利用4#线升温、3#线停产期间组织试拆装，评估拆除难度及所需专用工具，'
    '形成书面评估报告；在确认方案可行后，于后续停产窗口期正式实施衬板更换。'
)
add_body(
    '实施要求：更换前逐件核对备件型号及材质，更换后按规定完成空载试运，确认无异常后恢复正常运行状态。'
)

add_h3('2. #2磨辊密封更换')
add_body(
    '主要措施：与磨盘衬板更换同步安排实施，减少单独停机次数。'
    '更换前加密润滑油检测频次，及时掌握磨损发展趋势。'
)
add_body(
    '实施要求：更换密封时同步清洗润滑管路及油箱并更换新油，'
    '密封安装严格按设备安装手册相关技术要求执行。'
)

# ——（二）年度大修 ——
add_h2('（二）结合年度大修推进系统更新')
add_body_ni('整改目的：对可在年度检修中统一规划的系统性更新项目，纳入大修计划统筹安排。')

add_h3('1. 选粉机整体更换')
add_body(
    '主要措施：选粉机整体备件已在库，整体更换纳入明年度大修计划，'
    '与动氧系统年度检修同步实施。同步开展现场动平衡研究，确保更换后转子恢复正常运行状态。'
)
add_body(
    '实施要求：大修前将选粉机运行电流、振动值及煤粉细度纳入例行监测，做好趋势跟踪。'
)

add_h3('2. 布袋收尘器滤袋更换')
add_body(
    '主要措施：立即组织对各仓室滤袋进行抽样检查，根据检查结果评估各仓滤袋整体状况，'
    '制定分仓或整体更换方案。同步检查库存储备滤袋状态，如已老化需重新采购。'
)
add_body(
    '实施要求：更换完成后做好喷吹系统调试，建立滤袋运行寿命档案。'
)

# ——（三）运行保障措施 ——
add_h2('（三）持续完善运行保障措施')
add_body_ni('整改目的：对不需要长时间停产即可实施的改进项目，纳入日常运维持续完善。')

add_h3('1. 密封风改造')
add_body(
    '主要措施：为#1磨辊单独加装密封风管路，直接从密封风机出口引支管至密封罩，确保风压风量满足设计要求。'
)
add_body(
    '实施要求：改造后连续跟踪润滑油质变化，评估改造效果；若效果有限，再研究组合式密封方案。'
)

add_h3('2. 火花捕捉装置加装')
add_body(
    '主要措施：在热风炉出口至磨机入口管道上加装火花补给器，'
    '作为防止明火进入磨机系统的物理屏障，选型应结合系统工况确定。'
)
add_body(
    '实施要求：配备差压检测及报警装置，纳入定期检查维护。'
)

add_h3('3. 系统漏风治理')
add_body(
    '主要措施：全面排查制粉系统各法兰连接处、人孔门、检查孔等潜在漏风点并实施密封处理，'
    '在保证正常通风需求的前提下尽可能降低漏风率。'
)
add_body(
    '实施要求：漏风治理纳入日常巡检范围，发现漏风点及时处理。'
)

add_h3('4. 在线监测与联锁完善')
add_body(
    '主要措施：加强氧含量分析仪校准维护，完善CO浓度监测与磨机联锁保护逻辑，'
    '确保CO值达到报警限值时及时报警，达到上限时联锁停机并充入惰性气体。'
    '优化氮气惰化方案，确保紧急状态下能快速向磨机、布袋收尘器和粉煤仓注入足量氮气。'
)
add_body(
    '实施要求：定期检查联锁保护装置功能，确保有效可靠。'
)

add_h3('5. 运行管理标准化')
add_body(
    '主要措施：编制高挥发份煤种制粉操作指导书，明确不同煤种配比下的关键控制参数及操作原则。'
    '建立高挥发份煤种运行台账，每日记录关键参数，定期分析趋势及时调整。'
)
add_body(
    '实施要求：操作人员经专项培训后方可上岗。'
)

# ═══════════════════════════════════════════
# 三、整改计划
# ═══════════════════════════════════════════
add_h1('三、整改计划')
add_blank()

table = doc.add_table(rows=6, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(2.5), Cm(4.5), Cm(4.0), Cm(2.5)]
for row in table.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]

headers = ['整改项目', '主要内容', '计划时间', '责任单位']
data = [
    ['磨盘衬板\n更换',
     '完成磨盘衬板试拆装\n评估及正式更换',
     '试拆装：7月中旬\n正式更换：结合停产窗口',
     '冶炼部\n生产保障部'],
    ['#2磨辊密封\n更换',
     '完成磨辊密封组件\n更换及润滑系统清洁',
     '结合磨盘衬板更换\n同步实施',
     '冶炼部\n生产保障部'],
    ['选粉机整体\n更换',
     '完成选粉机整体更换\n并恢复转子动平衡',
     '纳入明年年度大修\n统筹安排',
     '冶炼部\n设备管理'],
    ['布袋收尘器\n滤袋抽检及更换',
     '完成滤袋抽样检查并\n制定更换实施方案',
     '抽检评估：立即开展\n分仓更换：逐步实施',
     '冶炼部\n设备管理'],
    ['高挥发份煤种\n安全保障措施',
     '完成火花捕捉器安装、\n漏风治理、监测联锁\n完善及操作规范编制',
     '火花捕捉器：3个月\n内完成\n其他措施：持续完善',
     '冶炼部\n设备管理\n供应部'],
]

for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    set_cell(cell, h, '黑体', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(cell, 'D9E2F3')

for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        al = WD_ALIGN_PARAGRAPH.CENTER if j in (2, 3) else None
        set_cell(cell, val, '仿宋', 12, False, al)

add_blank()

# ═══════════════════════════════════════════
# 四、下一步工作
# ═══════════════════════════════════════════
add_h1('四、下一步工作')
add_body(
    '下一步，冶炼部将持续加强立式磨关键部件状态监测，完善设备点检、润滑管理、在线监测及'
    '检修评估机制，及时掌握各部件运行状态，做到隐患早发现、早预警、早处置。'
)
add_body(
    '同时，结合年度检修及生产组织安排，按照"统筹安排、分步实施、风险可控"的原则，'
    '统筹推进设备更新改造及安全保障措施落实，持续提高立式磨系统安全性、可靠性和运行效率，'
    '为冶炼系统连续稳定生产提供设备保障。'
)

# 结尾
add_blank()
add_blank()
add_blank()
add_body_ni('妥否，请批示。')
add_blank()
add_blank()
add_right('印尼金川WP公司冶炼部')
add_right('2026年7月6日')

# 保存
output = r'D:\second-brain\submissions\立式磨存在问题及整改措施报告（正式上报版）.docx'
doc.save(output)
print(f'OK: {output}')
