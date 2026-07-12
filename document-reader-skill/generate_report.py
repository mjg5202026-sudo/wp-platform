"""Generate the vertical mill problem report as a Word document."""
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 全局样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '仿宋'
font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

# ── 辅助函数 ──
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(14)
    return h

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '仿宋'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(14)
    if bold:
        run.bold = True
    return p

# ════════════════════════════════════════════════════════════
# 正文
# ════════════════════════════════════════════════════════════

# ── 标题 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)
run = title.add_run('印尼金川WP公司立式磨存在问题及整改措施')
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(18)
run.bold = True

# ── 副标题 ──
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run('（冶炼部设备管理）')
run.font.name = '仿宋'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
run.font.size = Pt(14)

# ── 前言 ──
add_heading_styled('前  言', level=1)
add_para(
    '我公司粉煤制备系统采用布莱克散料处理技术（北京）有限公司生产的BRM28.3M立式辊磨机，'
    '自2019年投产运行至今已逾六年。该磨机设计产能55t/h（干粉），系统风量240,000m³/h，'
    '承担着为冶炼生产提供合格煤粉的关键任务。',
    indent=True
)
add_para(
    '随着运行时间的持续累积，立式磨各主要部件陆续出现不同程度的老化与磨损。近期排查发现，'
    '磨盘衬板、选粉机动叶片、布袋收尘器布袋、磨辊润滑密封系统等方面存在不同程度的安全隐患，'
    '此外，随着高挥发份煤种的应用趋于常态化，制粉系统的防爆安全要求也进一步提升。'
    '为保障设备安全、稳定、高效运行，现针对上述问题逐项分析并提出整改措施如下。',
    indent=True
)

# ════════════════════════════════════════════════════════════
# 一、磨盘衬板磨损
# ════════════════════════════════════════════════════════════
add_heading_styled('一、磨盘衬板磨损', level=1)

add_heading_styled('1.1 问题描述', level=2)
add_para(
    '立式磨磨盘衬板自投产以来已连续运行超过六年，长期承受物料碾压与研磨，工作面磨损严重，'
    '已接近或超过设计允许的磨损极限。磨盘衬板磨损直接导致以下问题：',
    indent=True
)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('研磨效率下降，磨机产能降低，单位电耗上升；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('料层控制难度增大，磨机运行稳定性变差；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('若继续超限使用，可能导致衬板断裂甚至磨盘本体损伤，造成更大的设备损失。')

add_heading_styled('1.2 整改措施', level=2)
add_para('（1）近期利用4#线升温、3#线停产期间（计划7月15日左右），组织检修人员对磨盘衬板进行试拆装，'
         '确认螺栓锈蚀程度、拆卸难度及所需专用工具，完整评估更换作业的工时与风险，形成书面拆装评估报告。', indent=True)
add_para('（2）在完成试拆装评估的基础上，编制详细的磨盘衬板更换施工方案，明确施工步骤、'
         '人员分工、吊装方案、安全措施及应急预案，方案经审核批准后实施。', indent=True)
add_para('（3）待3#线升温、1#线停产时，正式安排衬板更换作业。更换前确保备件已全部到场，'
         '并逐件核对衬板型号、尺寸与材质，确认与BRM28.3M磨机图纸（图号BRM28.3M.9）一致。', indent=True)
add_para('（4）更换完成后，进行空载试运行不少于30分钟，检查磨盘运转平稳性及衬板紧固情况，'
         '确认无异常后带料逐步加载至正常运行参数。', indent=True)

# ════════════════════════════════════════════════════════════
# 二、选粉机动叶片磨损
# ════════════════════════════════════════════════════════════
add_heading_styled('二、选粉机动叶片磨损', level=1)

add_heading_styled('2.1 问题描述', level=2)
add_para(
    '选粉机是控制煤粉细度的关键设备，其动叶片长期处于高速旋转及含尘气流冲刷工况下，'
    '叶片迎风面磨损明显，已影响分级效率。当前选粉机运行频率持续维持47Hz较高水平（据了解正常带载运行频率应在30-40Hz之间），'
    '反映出由于叶片磨损导致的分级效能下降，系统被迫以更高转速运行来达到细度要求。'
    '叶片磨损如持续发展，可能引发以下问题：',
    indent=True
)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('煤粉细度波动，不合格煤粉进入燃烧系统，影响回转窑工况；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('长时间高转速运行加速轴承及传动系统疲劳，缩短设备寿命；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('磨损加重后可能发生叶片脱落，造成选粉机本体损坏的重大设备事故。')

add_heading_styled('2.2 整改措施', level=2)
add_para('（1）鉴于选粉机整体有备件，且当前叶片磨损程度尚不完全影响运行，计划将选粉机整体更换纳入明年度大修计划，'
         '与动氧系统年度检修同步实施，充分利用停产窗口期，减少对生产的影响。', indent=True)
add_para('（2）在大修前，开展叶片更换及现场动平衡可行性的技术研究。建议与设备厂家或专业技术单位合作，'
         '研究在不整体拆下选粉机的情况下，实施叶片在线更换及现场动平衡校正的方案。'
         '若方案可行，可将更换作业时间从数周缩短至数天，极大降低停机损失。', indent=True)
add_para('（3）加强运行监控：每周对选粉机运行电流、振动值、转速及煤粉细度进行记录，'
         '建立趋势分析台账。一旦发现振动值突变或电流异常波动，立即停机检查。', indent=True)
add_para('（4）利用现有停产机会，定期检查选粉机静叶片安装角度是否与图纸一致，'
         '必要时使用专用工具进行调节，保证粗粉分离效能。', indent=True)

# ════════════════════════════════════════════════════════════
# 三、布袋收尘器布袋
# ════════════════════════════════════════════════════════════
add_heading_styled('三、布袋收尘器布袋老化', level=1)

add_heading_styled('3.1 问题描述', level=2)
add_para(
    '布袋收尘器是粉煤制备系统的重要安全与环保设备，其布袋自2021年12月更换至今，已连续运行近四年。'
    '聚酯针刺毡滤袋在长期处于含煤粉烟气、脉冲喷吹、温度波动等工况下，存在以下隐患：',
    indent=True
)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('滤袋纤维疲劳、强度下降，可能出现局部破损，导致粉尘排放超标，环保风险增加；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('滤袋透气性下降，系统阻力上升，引风机能耗增加，影响磨机通风量；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('破损布袋如未及时发现，煤粉可能进入净气室，在特定条件下构成安全风险。')

add_heading_styled('3.2 整改措施', level=2)
add_para('（1）立即组织对布袋收尘器各仓室进行抽样检查，按每个仓室抽取不少于3条布袋的原则，'
         '逐条检查布袋的表面磨损、破裂、硬化、烧灼痕迹等情况，拍照记录并建立检查台账。'
         '重点检查布袋口部、底部折边处及与笼骨接触部位。', indent=True)
add_para('（2）根据抽样检查结果，评估各仓布袋的整体状况，编制分仓更换计划。'
         '若抽查发现普遍老化严重（强度下降超过30%或存在多处破损），建议全部更换。'
         '若整体状况尚可，则按"先坏先换、分仓分批"的原则逐步更换。', indent=True)
add_para('（3）留存备件自采购至今已存放四年，需对库存布袋开箱检查是否存在受潮、发霉、'
         '老化脆变等问题。如库存布袋已严重老化不宜使用，需重新申报采购。'
         '检查结果记录至备件台账。', indent=True)
add_para('（4）更换完成后，做好新布袋的喷吹系统参数设定与调试。'
         '关注压差变化趋势，建立布袋运行寿命档案，明确下次更换的时间窗口建议。', indent=True)

# ════════════════════════════════════════════════════════════
# 四、磨辊润滑密封系统
# ════════════════════════════════════════════════════════════
add_heading_styled('四、磨辊润滑密封系统问题', level=1)

add_heading_styled('4.1 #2磨辊密封磨损', level=2)
add_para('存在问题：#2磨辊密封自投产至今一直未更换，密封件已严重磨损，'
         '煤粉颗粒已进入润滑系统，回油中可见煤粉。目前进回油尚正常，润滑功能基本正常，'
         '但若不及时处理，密封失效将进一步加剧，最终导致润滑系统损坏、磨辊轴承烧毁。', indent=True)
add_heading_styled('整改措施：', level=2)
add_para('（1）根据运行监控情况，择机整体更换#2磨辊密封组件，此项工作可与磨盘衬板更换同步安排，'
         '以减少单独停机的次数。', indent=True)
add_para('（2）更换前加强润滑油的定期取样检测，缩短取样周期为每两周一次，监测油液中铁谱、'
         '光谱及颗粒度指标，掌握磨损发展趋势。发现指标异常升高时立即反馈并提前安排检修。', indent=True)
add_para('（3）更换密封时，同步清洗润滑管路及油箱，更换润滑油，确保换油后润滑系统内无残留煤粉。'
         '安装密封件时严格按照BRM28.3M立式辊磨机安装手册相关要求执行。', indent=True)

add_heading_styled('4.2 #1磨辊密封风不足', level=2)
add_para('存在问题：#1磨辊位于磨辊密封风管道的尾端，管路阻力大，风压、风量不足，'
         '无法在密封罩内形成有效的正压气幕，导致煤粉进入密封罩内，逐步磨损密封件，'
         '最终威胁润滑系统。该问题属系统设计层面的管路布置缺陷——末端风压衰减所致。', indent=True)
add_heading_styled('整改措施：', level=2)
add_para('（1）近期应急方案：为#1磨辊单独加装一路密封风管路，直接从密封风机出口引支管至#1磨辊密封罩，'
         '确保密封风量和风压满足设计要求。安装后测试密封罩内静压，确保正压值达到设计指标。', indent=True)
add_para('（2）观察运行效果：加装后连续跟踪不少于一个月，定期检查#1磨辊润滑油质，'
         '评估单独供风方案的密封效果。若效果良好，将其纳入正式改造方案并固化管路走向。', indent=True)
add_para('（3）如若单独加装密封风方案效果不佳，需研究更换密封方式的技术方案，'
         '例如改为接触式密封或组合式密封（气封+接触密封），从根本上解决密封风不足的问题。'
         '此方案需与设备厂家或密封专业单位进行技术交流后确定。', indent=True)

# ════════════════════════════════════════════════════════════
# 五、高挥发份煤种防爆
# ════════════════════════════════════════════════════════════
add_heading_styled('五、高挥发份煤种应用安全措施', level=1)

add_heading_styled('5.1 问题描述', level=2)
add_para(
    '为降低生产成本、提升经济效益，目前采购的煤种中高挥发份烟煤配比逐渐提高。'
    '高挥发份煤种（Vdaf > 28%）在制粉过程中释放大量可燃挥发份，使制粉系统爆炸风险显著上升。'
    '立式磨制粉系统防爆的关键控制指标已在相关规程中明确规定：',
    indent=True
)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('当原料为烟煤时，制粉系统气粉混合物中O₂含量应控制≤14%；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('当原料为褐煤时，制粉系统气粉混合物中O₂含量应控制≤12%；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('一般认为烟气中O₂体积含量＞16%时易引起爆炸，＜14%时视为惰性气体状态。')
add_para(
    '目前烟气含氧量控制指标为≤12%（按褐煤标准从严执行），CO浓度控制≤500PPM。'
    '但在实际操作中，由于系统存在漏风，且热风炉运行中可能产生火花，'
    '制粉系统的防爆形势不容乐观。',
    indent=True
)

add_heading_styled('5.2 整改措施', level=2)
add_para('（1）在热风炉出口至磨机入口之间的管道上加装火花补给器（火花捕捉器），'
         '作为防止明火进入磨机系统的最后一道物理屏障。火花补给器的选型应考虑以下要点：', indent=True)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('安装位置：热风炉出口后段直管段，距热风炉出口不少于3倍管径处；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('设计流速：与管道内实际烟气流速匹配，保证捕捉效率不低于95%；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('材质要求：耐温不低于300℃，耐腐蚀，便于清灰维护；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('配备差压检测及报警装置，当火花器堵塞时及时提示清理。')

add_para('（2）系统漏风治理：全面排查制粉系统各法兰连接处、人孔门、检查孔、'
         '锁风阀等可能漏风点，进行密封处理。在保证系统正常通风需求的前提下，'
         '将漏风率控制在最低水平，抑制外部空气（O₂含量21%）进入系统导致含氧量上升。', indent=True)

add_para('（3）完善气体监测与联锁保护：', indent=True)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('加强在线氧含量分析仪的校准维护，确保数据准确可靠，每月至少校准一次；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('完善CO监测与磨机联锁保护逻辑，当CO浓度≥800PPM时自动触发报警，'
           '≥1200PPM时应联锁停止磨机并充入惰性气体；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('利用现有氮气系统，优化惰化方案，确保在紧急状态下能快速向磨机、'
           '布袋收尘器和粉煤仓注入足量氮气。')

add_para('（4）操作管理优化：', indent=True)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('编制高挥发份煤种制粉操作指导书，明确不同煤种配比下的目标氧含量、'
           '进出口温度控制范围、加料速率调整原则；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('操作人员需经专项培训并考核合格后方可上岗操作；')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('建立高挥发份煤种运行台账，记录每日煤种配比、氧含量、温度、CO等参数，'
           '定期分析趋势，及时调整。')

# ════════════════════════════════════════════════════════════
# 六、总结
# ════════════════════════════════════════════════════════════
add_heading_styled('六、总结与实施路线图', level=1)
add_para(
    '上述五项问题是当前立式磨系统面临的突出隐患，涉及磨机本体、分级系统、收尘系统、'
    '润滑密封系统及安全控制系统。根据问题的轻重缓急，建议按以下路线图推进整改：',
    indent=True
)

# 路线图表格
table = doc.add_table(rows=6, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['序号', '整改内容', '时间安排', '责任单位']
data = [
    ['1', '磨盘衬板试拆装评估', '7月15日左右（利用3#线停产窗口）', '冶炼部/生产保障部'],
    ['2', '布袋收尘器布袋抽样检查', '立即开展，2周内完成', '冶炼部/设备管理'],
    ['3', '#1磨辊单独加装密封风管路', '1个月内完成设计及安装', '冶炼部/生产保障部'],
    ['4', '火花补给器选型、采购及安装', '3个月内完成', '设备管理/供应部'],
    ['5', '磨盘衬板正式更换/\n#2磨辊密封更换/\n选粉机大修', '结合各线停产窗口\n分步实施，按明年度\n大修计划统筹安排', '冶炼部/生产保障部'],
]

for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(12)
    set_cell_shading(cell, 'D9E2F3')

for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '仿宋'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                run.font.size = Pt(12)

# ── 结尾 ──
add_para('')
add_para('请各相关部门按上述整改措施要求，结合生产计划统筹安排，确保立式磨系统安全、稳定、经济运行。', indent=True)

add_para('')
add_para('')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.add_run('印尼金川WP公司冶炼部').bold = True

p = doc.add_paragraph()
p.add_run('2025年7月')

# ── 保存 ──
output_path = r'D:\工作\设备资料\分类资料\立式磨\印尼金川WP公司立式磨存在问题及整改措施报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
