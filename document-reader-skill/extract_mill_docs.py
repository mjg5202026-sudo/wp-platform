"""Extract text from key vertical mill documents for report writing."""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

files = [
    r"D:\工作\设备资料\分类资料\立式磨\立式磨设备大修检修技术方案 - 20240421.docx",
    r"D:\工作\设备资料\分类资料\立式磨\立式磨衬板更换检修方案.docx",
    r"D:\工作\设备资料\分类资料\立式磨\立式磨资料\会议纪要\关于立式磨油路系统等存在问题专题会议纪要（2019—001期） 20190417.docx",
    r"D:\工作\设备资料\分类资料\立式磨\立式磨资料\会议纪要\立式磨技术讨论会 20190721.docx",
    r"D:\工作\设备资料\分类资料\立式磨\立式磨资料\立式磨改动方案\BRM28.3M立式辊磨机调试方法和故障处理方法.docx",
    r"D:\工作\设备资料\分类资料\立式磨\粉煤制备系统(考拉海).docx",
]

for filepath in files:
    p = Path(filepath)
    if not p.exists():
        print(f"\n=== SKIP (not found): {p.name} ===")
        continue

    print(f"\n{'='*60}")
    print(f"=== {p.name} ===")
    print(f"{'='*60}")

    doc = Document(str(p))

    # Print paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            print(text)

    # Print tables
    for ti, table in enumerate(doc.tables):
        print(f"\n[TABLE {ti+1}]")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(" | ".join(cells))
